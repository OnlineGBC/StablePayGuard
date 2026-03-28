"""Generate StablePayGuard MANUAL.docx from structured content."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ---------------------------------------------------------------------------
# Styles
# ---------------------------------------------------------------------------

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def set_heading(paragraph, level, text):
    paragraph.text = text
    if level == 0:
        paragraph.style = doc.styles['Title']
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        paragraph.style = doc.styles[f'Heading {level}']

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    return p

def add_paragraph(doc, text=''):
    return doc.add_paragraph(text)

def add_code(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles['No Spacing']
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1F, 0x69, 0x35)
    # light grey background via paragraph shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F2F2F2')
    pPr.append(shd)
    return p

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    # header row
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
        hdr[i].paragraphs[0].runs[0].bold = True if hdr[i].paragraphs[0].runs else None
        # bold manually
        p = hdr[i].paragraphs[0]
        run = p.runs[0] if p.runs else p.add_run(h)
        run.bold = True
        if not p.runs:
            hdr[i].text = ''
            run2 = hdr[i].paragraphs[0].add_run(h)
            run2.bold = True
    # data rows
    for r, row in enumerate(rows):
        cells = table.rows[r + 1].cells
        for c, val in enumerate(row):
            cells[c].text = val
    doc.add_paragraph()
    return table

# ---------------------------------------------------------------------------
# Title Page
# ---------------------------------------------------------------------------

title = doc.add_heading('StablePayGuard', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub = doc.add_paragraph('Complete User and Technical Manual')
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].font.size = Pt(14)
sub.runs[0].bold = True

doc.add_paragraph()
ver = doc.add_paragraph('Version 1.0  |  March 2026')
ver.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# ---------------------------------------------------------------------------
# Table of Contents (manual — Word will update on open)
# ---------------------------------------------------------------------------

add_heading(doc, 'Table of Contents', level=1)
toc_entries = [
    ('1.', 'What Is StablePayGuard?'),
    ('2.', 'Who Benefits and How'),
    ('3.', 'Where the Code Lives'),
    ('4.', 'Prerequisites'),
    ('5.', 'Running Locally'),
    ('6.', 'Accessing the Live App'),
    ('7.', 'Usage Guide'),
    ('   7.1', 'Logging In'),
    ('   7.2', 'Dashboard'),
    ('   7.3', 'Creating a Policy'),
    ('   7.4', 'Submitting a Payment Intent'),
    ('   7.5', 'Connecting a Wallet'),
    ('   7.6', 'Viewing Policies'),
    ('8.', 'Demo Mode vs. Live Blockchain Mode'),
    ('9.', 'Environment Variables'),
    ('10.', 'Changing the Admin Password'),
    ('11.', 'Connecting a Real AI Key'),
    ('12.', 'Infrastructure on Google Cloud'),
    ('   12.1', 'GCP Project'),
    ('   12.2', 'Cloud Run'),
    ('   12.3', 'Cloud SQL'),
    ('   12.4', 'Google Secret Manager'),
    ('   12.5', 'APIs Enabled'),
    ('13.', 'Redeploying the App'),
    ('14.', 'Viewing Logs'),
    ('15.', 'Codebase Overview'),
    ('   15.1', 'Directory Structure'),
    ('   15.2', 'Backend Components'),
    ('   15.3', 'Frontend'),
    ('   15.4', 'Smart Contract'),
    ('   15.5', 'Tests'),
    ('   15.6', 'CI/CD'),
    ('16.', 'Going Live on the Blockchain'),
    ('17.', 'Security'),
    ('18.', 'Troubleshooting'),
    ('19.', 'Known Limitations'),
    ('20.', 'Full API Reference'),
    ('21.', 'Architecture Deep Dive'),
    ('   21.1', 'System Diagram'),
    ('   21.2', 'Component Details'),
    ('   21.3', 'Database Schema'),
    ('   21.4', 'Data Flow: Payment Execution'),
    ('   21.5', 'Data Flow: Secret Loading'),
    ('   21.6', 'Deployment Architecture'),
    ('22.', 'Roadmap'),
    ('23.', 'Business Scenario'),
    ('   23.1', 'Ideal Customer Profile'),
    ('   23.2', 'The Core Problem This Solves'),
    ('   23.3', 'Why This Needs Crypto'),
    ('   23.4', 'Cross-Border Payment Economics'),
    ('   23.5', 'Agent A: SaaS Vendor Payments'),
    ('   23.6', 'Agent B: Global Contractor Payments'),
    ('   23.7', 'Agent C: Employee T&E Reimbursements'),
    ('   23.8', 'All Three Agents Together'),
    ('   23.9', 'Traditional AP vs. StablePayGuard'),
    ('   23.10', 'Security Considerations'),
    ('   23.11', 'The Fundamental Principle'),
    ('   23.12', 'Production Roadmap'),
    ('24.', 'Contributing Guidelines'),
]
for num, title_text in toc_entries:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3) if num.startswith(' ') else Inches(0)
    run = p.add_run(f'{num.strip()}  {title_text}')
    run.font.size = Pt(11)

doc.add_page_break()

# ---------------------------------------------------------------------------
# Section 1
# ---------------------------------------------------------------------------

add_heading(doc, '1. What Is StablePayGuard?', level=1)
add_paragraph(doc, (
    'StablePayGuard is a web-based control panel for managing and enforcing AI-driven '
    'payment rules. In plain terms:'
))
quote = doc.add_paragraph(
    'Imagine you have an AI assistant that is allowed to make payments on your behalf — '
    'paying cloud hosting bills, vendor invoices, software subscriptions. Without controls '
    'in place, the AI could overspend, pay the wrong recipient, or act outside its '
    'authorised scope. StablePayGuard solves this by letting you define strict spending '
    'rules ("policies") that every payment must pass before it goes through.'
)
quote.paragraph_format.left_indent = Inches(0.4)
quote.paragraph_format.right_indent = Inches(0.4)
quote.runs[0].italic = True

add_paragraph(doc, 'A policy answers questions like:')
for item in [
    'Which AI agent is allowed to spend?',
    'In which token (e.g. USDC, ETH, DAI)?',
    'What is the total budget for this agent?',
    'What is the maximum it can spend in a single transaction?',
    'What is the stated purpose (e.g. "cloud infrastructure", "vendor payments")?',
]:
    p = doc.add_paragraph(item, style='List Bullet')

add_paragraph(doc, (
    'When a payment is attempted, it is checked against the matching policy. If it exceeds '
    'any limit, it is rejected automatically.'
))
add_paragraph(doc, (
    'StablePayGuard also includes an AI payment intent parser: you type a natural-language '
    'instruction such as "Pay AWS $200 for cloud hosting" and the system converts it into '
    'a structured payment record — recipient, amount, token, purpose — ready for policy enforcement.'
))

# ---------------------------------------------------------------------------
# Section 2
# ---------------------------------------------------------------------------

add_heading(doc, '2. Who Benefits and How', level=1)
add_paragraph(doc, (
    'Any organisation that uses AI to automate payments, procurement, or financial operations '
    'benefits from having a policy layer. StablePayGuard gives finance and ops teams visibility '
    'and control without needing to understand the underlying AI code.'
))
add_table(doc,
    ['Participant', 'Role', 'Benefit'],
    [
        ['Finance / Ops Manager', 'Sets spending policies', 'Ensures AI agents never exceed budget or act outside authorised scope'],
        ['AI Agent', 'Submits payment intents', 'Gets a clear approved/rejected decision, enabling automated retry or escalation'],
        ['Developer / Admin', 'Deploys and maintains the system', 'Structured API, persistent database, and audit trail rather than ad-hoc scripts'],
        ['Auditor / Compliance', 'Reviews transaction history', 'Full log of every payment attempt, approval, and rejection with timestamps'],
        ['Blockchain / DeFi Operator', 'Uses on-chain enforcement', 'Policies optionally enforced on Ethereum (Sepolia testnet), making them tamper-proof'],
    ]
)

# ---------------------------------------------------------------------------
# Section 3
# ---------------------------------------------------------------------------

add_heading(doc, '3. Where the Code Lives', level=1)

add_heading(doc, 'Local Machine', level=2)
add_paragraph(doc, 'The source code lives at:')
add_code(doc, 'C:\\Users\\raja\\StablePayGuard\\')
add_paragraph(doc, 'This is a Git repository. All code changes are committed and pushed from here.')

add_heading(doc, 'GitHub', level=2)
add_paragraph(doc, 'The repository is hosted at:')
add_code(doc, 'https://github.com/OnlineGBC/StablePayGuard')
add_paragraph(doc, (
    'This is the source of truth for all deployments. Cloud builds pull from a local push '
    'via gcloud run deploy --source . which uploads sources directly to Google Cloud Build.'
))

add_heading(doc, 'Google Cloud — Project: stablepayguard', level=2)
add_table(doc,
    ['Component', 'Location'],
    [
        ['Cloud Run service', 'us-east1 region, service name: stablepayguard'],
        ['Container image', 'Artifact Registry: us-east1-docker.pkg.dev/stablepayguard/cloud-run-source-deploy/stablepayguard'],
        ['Cloud SQL database', 'Instance: stablepayguard-db, Postgres 15, region us-east1'],
        ['Database name', 'stablepayguard'],
        ['Live URL', 'https://stablepayguard-684704256193.us-east1.run.app'],
    ]
)

# ---------------------------------------------------------------------------
# Section 4
# ---------------------------------------------------------------------------

add_heading(doc, '4. Prerequisites', level=1)

add_heading(doc, 'To Run Locally', level=2)
add_table(doc,
    ['Requirement', 'Purpose'],
    [
        ['Python 3.10+', 'Runs the Flask application'],
        ['pip', 'Installs Python dependencies'],
        ['Git', 'Clones the repository'],
        ['PostgreSQL 15+ or Docker (optional)', 'Production database. If DATABASE_URL is not set, SQLite is used automatically for local development.'],
    ]
)

add_heading(doc, 'To Deploy to Google Cloud', level=2)
add_table(doc,
    ['Requirement', 'Purpose'],
    [
        ['Google Cloud SDK (gcloud)', 'Deploys to Cloud Run, manages Cloud SQL'],
        ['A GCP account with billing enabled', 'Cloud Run and Cloud SQL incur costs'],
        ['Docker (optional)', 'Only needed to build/test the container locally'],
    ]
)

add_heading(doc, 'Optional (for full functionality)', level=2)
add_table(doc,
    ['Requirement', 'Purpose'],
    [
        ['Anthropic API key (SYNTH_API_KEY)', 'Powers the AI payment intent parser'],
        ['OpenAI API key (OPENAI_API_KEY)', 'Fallback AI provider'],
        ['Infura / Alchemy RPC URL', 'Connects to Ethereum blockchain'],
        ['Ethereum private key', 'Signs on-chain policy transactions'],
    ]
)

# ---------------------------------------------------------------------------
# Section 5
# ---------------------------------------------------------------------------

add_heading(doc, '5. Running Locally', level=1)

add_heading(doc, 'Option A — Docker Compose (Recommended)', level=2)
add_paragraph(doc, '1. Clone the repository:')
add_code(doc, 'git clone https://github.com/OnlineGBC/StablePayGuard\ncd StablePayGuard')
add_paragraph(doc, '2. Copy the example environment file and fill in your values:')
add_code(doc, 'cp .env.example .env')
add_paragraph(doc, '3. Start the database and web service:')
add_code(doc, 'docker-compose up')
add_paragraph(doc, '4. Open http://localhost:5000 in your browser.')
add_paragraph(doc, (
    'Docker Compose starts a Postgres 15 container (stablepayguard-db) and the Flask web '
    'container (stablepayguard). Data persists in a Docker volume (postgres_data) across restarts.'
))

add_heading(doc, 'Option B — SQLite (Simplest, No Database Required)', level=2)
add_paragraph(doc, 'If DATABASE_URL is not set (or is commented out in .env), the app automatically uses SQLite. No PostgreSQL or Docker needed.')
add_paragraph(doc, '1. Install dependencies:')
add_code(doc, 'pip install -r requirements.txt')
add_paragraph(doc, '2. Create a .env file and comment out DATABASE_URL:')
add_code(doc, 'cp .env.example .env\n# Edit .env: comment out DATABASE_URL line')
add_paragraph(doc, '3. Run the app:')
add_code(doc, 'python app/app.py')
add_paragraph(doc, '4. Open http://localhost:5000. A SQLite database file is created automatically in app/instance/.')

add_heading(doc, 'Option C — Manual PostgreSQL (No Docker)', level=2)
add_paragraph(doc, '1. Start a local PostgreSQL instance and create the database and user:')
add_code(doc, "CREATE USER stablepayguard WITH PASSWORD 'stablepayguard';\nCREATE DATABASE stablepayguard OWNER stablepayguard;")
add_paragraph(doc, '2. Install Python dependencies:')
add_code(doc, 'pip install -r requirements.txt')
add_paragraph(doc, '3. Create a .env file:')
add_code(doc, """\
# AI providers (at least one recommended for live payment intent parsing)
SYNTH_API_KEY=<your-anthropic-api-key>
OPENAI_API_KEY=<your-openai-api-key>

# Flask
FLASK_ENV=development
PORT=5000
SECRET_KEY=<any-long-random-string>
ADMIN_PASSWORD=<your-chosen-password>

# Database (comment out to use SQLite automatically)
DATABASE_URL=postgresql://stablepayguard:stablepayguard@localhost:5432/stablepayguard

# Blockchain / Sepolia testnet (all optional — omit for demo mode)
RPC_URL=https://sepolia.infura.io/v3/<your-infura-key>
PRIVATE_KEY=<your-wallet-private-key>
POLICY_CONTRACT=0x16229C14aAa18C7bC069f5b9092f5Af8884f3781
OWNER_WALLET=<your-wallet-address>""")
add_paragraph(doc, 'Note: The .env file is listed in .gitignore and will never be committed to the repository.')
add_paragraph(doc, '4. Run the app:')
add_code(doc, 'python app/app.py')
add_paragraph(doc, '5. Open http://localhost:5000.')

# ---------------------------------------------------------------------------
# Section 6
# ---------------------------------------------------------------------------

add_heading(doc, '6. Accessing the Live App', level=1)
add_paragraph(doc, 'The live deployment is at:')
add_code(doc, 'https://stablepayguard-684704256193.us-east1.run.app')
add_paragraph(doc, 'Login credentials:')
for item in ['Username: (none — single user system)', 'Password: demo']:
    doc.add_paragraph(item, style='List Bullet')
add_paragraph(doc, 'To change the password, see Section 10.')

# ---------------------------------------------------------------------------
# Section 7
# ---------------------------------------------------------------------------

add_heading(doc, '7. Usage Guide', level=1)

add_heading(doc, '7.1 Logging In', level=2)
add_paragraph(doc, '[SCREENSHOT: Login screen showing the StablePayGuard logo and password field]').runs[0].italic = True
add_paragraph(doc, (
    'When you open the app, a login overlay appears. Enter the admin password (demo by default) '
    'and click Sign In. The overlay disappears and the dashboard loads. '
    'If the password is wrong, a red error message appears below the input field.'
))

add_heading(doc, '7.2 Dashboard', level=2)
add_paragraph(doc, '[SCREENSHOT: Dashboard with KPI cards, activity feed, and transaction table]').runs[0].italic = True
add_paragraph(doc, 'The dashboard is the home screen. It shows:')
add_table(doc,
    ['Section', 'What it shows'],
    [
        ['KPI cards', 'Total policies, total payments, total payment volume (USD), approval rate (%)'],
        ['Payment Chart', 'Bar chart of payment volume by day of the week'],
        ['Recent Transactions', 'Last 10 transactions with status badges (Completed / Pending / Declined)'],
        ['Activity Feed', 'Last 10 system events (policy created, payment executed, wallet connected)'],
        ['Wallet status', 'Whether a wallet is connected and its address'],
    ]
)
add_paragraph(doc, 'All data refreshes when the page loads. To see updated numbers after an action, navigate away and back, or refresh the page.')

add_heading(doc, '7.3 Creating a Policy', level=2)
add_paragraph(doc, '[SCREENSHOT: Policy creation form with fields filled in]').runs[0].italic = True
add_paragraph(doc, '1. Click Policies in the left sidebar.\n2. Fill in the Create New Policy form:')
add_table(doc,
    ['Field', 'Description', 'Example'],
    [
        ['Agent Address', 'Ethereum address of the AI agent being authorised', '0xAgentABC123'],
        ['Token', 'Token the agent is allowed to spend', 'USDC'],
        ['Total Budget', 'Maximum lifetime spend for this policy', '5000'],
        ['Per-Tx Limit', 'Maximum spend per single transaction', '500'],
        ['Valid From', 'Unix timestamp when policy starts (0 = now)', '0'],
        ['Valid Until', 'Unix timestamp when policy expires (0 = no expiry)', '0'],
        ['Purpose', 'Human-readable description of payment purpose', 'vendor payments'],
    ]
)
add_paragraph(doc, '3. Click Create Policy.\n4. A success message shows the Policy ID and transaction hash.\n5. The Dashboard KPI counter for Policies increments by 1.')
add_paragraph(doc, 'Validation rules:')
for item in [
    'Agent address cannot be empty',
    'Total budget must be greater than 0',
    'Purpose cannot be empty',
    'Per-tx limit cannot exceed total budget',
]:
    doc.add_paragraph(item, style='List Bullet')
add_paragraph(doc, 'If any rule is violated, a 400 Bad Request response is returned with a field-level error message.')

add_heading(doc, '7.4 Deactivating a Policy', level=2)
add_paragraph(doc, '1. Click Policies in the sidebar.\n2. Select a policy from the list.\n3. Click Deactivate Policy.\n4. The policy is deactivated on-chain (or in demo mode). No further payments can be made against it.')
add_paragraph(doc, 'Deactivation is permanent. To re-enable spending, create a new policy.')

add_heading(doc, '7.5 Submitting a Payment Intent', level=2)
add_paragraph(doc, '[SCREENSHOT: Payment intent form with a natural-language task entered]').runs[0].italic = True
add_paragraph(doc, '1. Click Payments in the sidebar.\n2. In the Payment Intent panel, type a natural-language task, for example:')
for item in [
    '"Pay AWS $200 for cloud hosting"',
    '"Send $50 USDC to Stripe for subscription fees"',
    '"Transfer 100 DAI to vendor 0xABC for consulting"',
]:
    doc.add_paragraph(item, style='List Bullet')
add_paragraph(doc, '3. Click Generate Intent.\n4. The AI parses the text and returns a structured result:')
add_code(doc, '{\n  "recipient": "AWS",\n  "amount": 200,\n  "token": "USDC",\n  "purpose": "cloud hosting"\n}')
add_paragraph(doc, '5. The intent is stored in the database and appears in the transaction feed on the Dashboard.')
add_paragraph(doc, 'Note: In demo mode (no AI API key configured), a pre-built demo response is returned. See Section 11 to enable live AI parsing.')

add_heading(doc, '7.6 Executing a Payment', level=2)
add_paragraph(doc, '1. Click Payments in the sidebar.\n2. In the Execute Payment panel, fill in:')
add_table(doc,
    ['Field', 'Description', 'Example'],
    [
        ['Policy ID', 'The policy to charge this payment against', 'POL-101'],
        ['Recipient', 'Name or address of the payment recipient', 'AWS'],
        ['Amount', 'Payment amount in USD', '500'],
        ['Purpose', 'Optional description of the payment', 'cloud hosting'],
    ]
)
add_paragraph(doc, '3. Click Execute Payment.\n4. The system validates the amount against the policy\'s remaining budget. If it would exceed the budget, the request is rejected with a 422 error.\n5. If approved, the payment is submitted on-chain (live mode) or recorded as a demo transaction. The transaction appears in the ledger.')

add_heading(doc, '7.7 Viewing and Filtering Transactions', level=2)
add_paragraph(doc, '1. Click Transactions in the sidebar to see the full transaction ledger.\n2. Filter by status or policy using the query parameters:')
add_code(doc, 'GET /api/transactions?status=Completed&policy=POL-101&limit=20&offset=0')
add_paragraph(doc, 'Available status values: Completed, Pending, Declined.')

add_heading(doc, '7.8 Connecting a Wallet', level=2)
add_paragraph(doc, '1. Click the Connect Wallet button in the top bar, or navigate to the Wallet section.\n'
               '2. Optionally enter a valid Ethereum address (e.g. 0xd8dA6BF26964aF9D7eEd9e03E53415D37aA96045). If left blank, a demo address is used.\n'
               '3. Click Connect.\n'
               '4. The top bar updates to show the connected address.\n'
               '5. The Dashboard wallet panel shows connected: true.')
add_paragraph(doc, 'Validation: The address must be a valid Ethereum address (42 characters, starting with 0x, valid hex). Invalid addresses return a 400 error.')

add_heading(doc, '7.9 Viewing Policies', level=2)
add_paragraph(doc, '1. Click Policies in the sidebar.\n2. The policy list shows all created policies with:')
for item in [
    'Policy ID', 'Agent address', 'Token', 'Total budget',
    'Remaining budget (total minus all completed transactions against this policy)',
    'Spent amount', 'Purpose',
]:
    doc.add_paragraph(item, style='List Bullet')
add_paragraph(doc, 'Remaining budget is calculated live from the transaction database each time the dashboard loads.')

# ---------------------------------------------------------------------------
# Section 8
# ---------------------------------------------------------------------------

add_heading(doc, '8. Demo Mode vs. Live Blockchain Mode', level=1)

add_heading(doc, 'Demo Mode (Default)', level=2)
for item in [
    'No Ethereum RPC connection required',
    'Policy creation returns a fake transaction hash prefixed with demo-',
    'Payment enforcement is simulated in the database only',
    'The /api/contract/status endpoint returns "mode": "demo"',
    'All data persists to PostgreSQL — policies, transactions, and activity logs all work normally',
]:
    doc.add_paragraph(item, style='List Bullet')

add_heading(doc, 'Live Blockchain Mode', level=2)
for item in [
    'Requires RPC_URL, PRIVATE_KEY, and POLICY_CONTRACT environment variables',
    'Policy creation submits a real transaction to the Ethereum blockchain (Sepolia testnet)',
    'The smart contract enforces per-transaction limits on-chain',
    'Transaction hash is a real Ethereum tx hash viewable on Etherscan',
    'The /api/contract/status endpoint returns "mode": "live"',
]:
    doc.add_paragraph(item, style='List Bullet')
add_paragraph(doc, 'To switch from demo to live, see Section 16.')

# ---------------------------------------------------------------------------
# Section 9
# ---------------------------------------------------------------------------

add_heading(doc, '9. Environment Variables', level=1)
add_paragraph(doc, 'All configuration is done via environment variables. For local development, put these in a .env file in the project root.')
add_table(doc,
    ['Variable', 'Required', 'Default', 'Description'],
    [
        ['DATABASE_URL', 'Yes', '—', 'PostgreSQL connection string. Format: postgresql://user:password@host:port/dbname'],
        ['SECRET_KEY', 'Yes', '—', 'Flask session secret. Use a long random string in production'],
        ['ADMIN_PASSWORD', 'Yes', '—', 'Password for the single admin login'],
        ['SYNTH_API_KEY', 'No', '—', 'Anthropic API key (primary AI provider for payment intent parsing)'],
        ['OPENAI_API_KEY', 'No', '—', 'OpenAI API key (fallback AI provider)'],
        ['RPC_URL', 'No', '—', 'Ethereum RPC endpoint (e.g. Infura Sepolia URL). Required for live blockchain mode'],
        ['PRIVATE_KEY', 'No', '—', 'Ethereum wallet private key for signing on-chain transactions'],
        ['POLICY_CONTRACT', 'No', '—', 'Deployed PolicyManager contract address'],
        ['OWNER_WALLET', 'No', '—', 'Ethereum address of the contract owner'],
        ['SEED_DATA', 'No', 'false', 'Set to true to seed demo transactions and policies on first startup'],
        ['PORT', 'No', '5000', 'Port the app listens on (Cloud Run sets this automatically to 8080)'],
        ['FLASK_ENV', 'No', 'production', 'Set to development for debug mode locally'],
    ]
)

# ---------------------------------------------------------------------------
# Section 10
# ---------------------------------------------------------------------------

add_heading(doc, '10. Changing the Admin Password', level=1)

add_heading(doc, 'On Cloud Run (via Secret Manager)', level=2)
add_paragraph(doc, 'Since ADMIN_PASSWORD is stored in Google Secret Manager, update the secret value rather than the env var:')
add_code(doc, 'echo -n "yournewpassword" | gcloud secrets versions add ADMIN_PASSWORD \\\n  --project stablepayguard \\\n  --data-file=-')
add_paragraph(doc, 'Then redeploy to pick up the new secret version:')
add_code(doc, 'gcloud run deploy stablepayguard \\\n  --source . \\\n  --region us-east1 \\\n  --project stablepayguard \\\n  --quiet')

add_heading(doc, 'Locally', level=2)
add_paragraph(doc, 'Update ADMIN_PASSWORD in your .env file and restart the app.')

# ---------------------------------------------------------------------------
# Section 11
# ---------------------------------------------------------------------------

add_heading(doc, '11. Connecting a Real AI Key', level=1)
add_paragraph(doc, 'Without an AI key, the payment intent parser returns a canned demo response. To enable live AI parsing:')

add_heading(doc, 'Using Anthropic (Recommended)', level=2)
add_paragraph(doc, '1. Obtain an API key from console.anthropic.com\n2. Set the environment variable:')
add_code(doc, 'gcloud run services update stablepayguard \\\n  --region us-east1 \\\n  --project stablepayguard \\\n  --update-env-vars "SYNTH_API_KEY=sk-ant-your-key-here"')
add_paragraph(doc, 'The app uses claude-haiku-4-5-20251001 for fast, cost-effective parsing.')

add_heading(doc, 'Using OpenAI (Fallback)', level=2)
add_code(doc, 'gcloud run services update stablepayguard \\\n  --region us-east1 \\\n  --project stablepayguard \\\n  --update-env-vars "OPENAI_API_KEY=sk-your-openai-key"')
add_paragraph(doc, 'The fallback chain is: SYNTH_API_KEY → OPENAI_API_KEY → demo response.')

# ---------------------------------------------------------------------------
# Section 12
# ---------------------------------------------------------------------------

add_heading(doc, '12. Infrastructure on Google Cloud', level=1)

add_heading(doc, 'GCP Project', level=2)
add_table(doc,
    ['Property', 'Value'],
    [
        ['Project ID', 'stablepayguard'],
        ['Project Name', 'StablePayGuard'],
        ['Billing account', '01D45F-38C7C5-24A487 (same as fednowrtppayrails)'],
        ['Region', 'us-east1'],
    ]
)

add_heading(doc, 'Cloud Run', level=2)
add_table(doc,
    ['Property', 'Value'],
    [
        ['Service name', 'stablepayguard'],
        ['URL', 'https://stablepayguard-684704256193.us-east1.run.app'],
        ['Concurrency', '8 threads per instance, 1 worker'],
        ['Authentication', 'Public (app-level auth via session cookie)'],
        ['Cloud SQL connection', 'Unix socket via stablepayguard:us-east1:stablepayguard-db'],
    ]
)

add_heading(doc, 'Cloud SQL', level=2)
add_table(doc,
    ['Property', 'Value'],
    [
        ['Instance name', 'stablepayguard-db'],
        ['Database version', 'PostgreSQL 15'],
        ['Tier', 'db-f1-micro (1 shared vCPU, 614 MB RAM)'],
        ['Region', 'us-east1'],
        ['Database', 'stablepayguard'],
        ['User', 'stablepayguard'],
        ['Connection name', 'stablepayguard:us-east1:stablepayguard-db'],
    ]
)

add_heading(doc, '12.4 Google Secret Manager', level=2)
add_paragraph(doc, (
    'All sensitive credentials are stored in Google Secret Manager rather than as plain '
    'environment variables. Cloud Run pulls them at runtime via the --update-secrets flag. '
    'This means secrets are never visible in the Cloud Console environment variables panel.'
))
add_paragraph(doc, 'The following secrets are currently configured in project stablepayguard:')
add_table(doc,
    ['Secret Name', 'Description'],
    [
        ['SECRET_KEY', 'Flask session signing key'],
        ['ADMIN_PASSWORD', 'Dashboard login password'],
        ['SYNTH_API_KEY', 'Anthropic API key for AI payment intent parsing'],
        ['OPENAI_API_KEY', 'OpenAI API key (fallback AI provider)'],
        ['RPC_URL', 'Infura Sepolia RPC endpoint for Ethereum blockchain'],
        ['PRIVATE_KEY', 'Ethereum wallet private key for signing on-chain transactions'],
        ['POLICY_CONTRACT', 'Deployed PolicyManager smart contract address (Sepolia)'],
        ['OWNER_WALLET', 'Ethereum wallet address of the contract owner'],
    ]
)
add_paragraph(doc, 'To add or update a secret:')
add_code(doc, '# Create a new secret\necho -n "value" | gcloud secrets create SECRET_NAME \\\n  --project stablepayguard \\\n  --replication-policy automatic \\\n  --data-file=-\n\n# Update an existing secret with a new version\necho -n "newvalue" | gcloud secrets versions add SECRET_NAME \\\n  --project stablepayguard \\\n  --data-file=-')
add_paragraph(doc, 'To view all secrets (names only — values are never shown):')
add_code(doc, 'gcloud secrets list --project stablepayguard')
add_paragraph(doc, 'To grant Cloud Run access to a new secret:')
add_code(doc, 'gcloud secrets add-iam-policy-binding SECRET_NAME \\\n  --project stablepayguard \\\n  --member="serviceAccount:684704256193-compute@developer.gserviceaccount.com" \\\n  --role="roles/secretmanager.secretAccessor"')

add_heading(doc, '12.5 APIs Enabled', level=2)
for api in ['Cloud Run API', 'Cloud SQL Admin API', 'Cloud Build API', 'Artifact Registry API', 'Secret Manager API']:
    doc.add_paragraph(api, style='List Bullet')

# ---------------------------------------------------------------------------
# Section 13
# ---------------------------------------------------------------------------

add_heading(doc, '13. Redeploying the App', level=1)
add_paragraph(doc, 'After making code changes locally:')
add_code(doc, 'gcloud run deploy stablepayguard \\\n  --source . \\\n  --region us-east1 \\\n  --project stablepayguard \\\n  --quiet')
add_paragraph(doc, 'This triggers Cloud Build to rebuild the Docker image and deploy a new revision. Existing environment variables are preserved unless you pass --set-env-vars or --update-env-vars.')

add_paragraph(doc, 'To update env vars without changing code:')
add_code(doc, 'gcloud run services update stablepayguard \\\n  --region us-east1 \\\n  --project stablepayguard \\\n  --update-env-vars "KEY=value"')

add_paragraph(doc, 'To roll back to a previous revision:')
add_code(doc, 'gcloud run services update-traffic stablepayguard \\\n  --region us-east1 \\\n  --project stablepayguard \\\n  --to-revisions REVISION_NAME=100')

add_paragraph(doc, 'To list available revisions:')
add_code(doc, 'gcloud run revisions list --service stablepayguard \\\n  --region us-east1 \\\n  --project stablepayguard')

# ---------------------------------------------------------------------------
# Section 14
# ---------------------------------------------------------------------------

add_heading(doc, '14. Viewing Logs', level=1)

add_heading(doc, 'All Logs', level=2)
add_code(doc, 'gcloud logging read \\\n  "resource.type=cloud_run_revision AND resource.labels.service_name=stablepayguard" \\\n  --project stablepayguard \\\n  --limit 50 \\\n  --format="value(timestamp,textPayload)"')

add_heading(doc, 'Errors Only', level=2)
add_code(doc, 'gcloud logging read \\\n  "resource.type=cloud_run_revision AND resource.labels.service_name=stablepayguard AND severity>=ERROR" \\\n  --project stablepayguard \\\n  --limit 20')

add_heading(doc, 'From the Google Cloud Console', level=2)
for item in [
    'Go to console.cloud.google.com',
    'Select project stablepayguard',
    'Navigate to Cloud Run → stablepayguard → Logs tab',
]:
    doc.add_paragraph(item, style='List Number')

# ---------------------------------------------------------------------------
# Section 15
# ---------------------------------------------------------------------------

add_heading(doc, '15. Codebase Overview', level=1)

add_heading(doc, '15.1 Directory Structure', level=2)
add_code(doc, """\
StablePayGuard/
├── app/
│   ├── app.py                  Flask app entry point, dashboard API, blueprint registration
│   ├── models.py               SQLAlchemy database models
│   ├── store.py                Database access layer (CRUD functions)
│   ├── schemas.py              Pydantic request validation schemas
│   ├── extensions.py           Flask-Limiter rate limiter instance
│   ├── utils.py                login_required decorator, validate_request helper
│   ├── routes/
│   │   ├── auth.py             POST /api/auth/login, logout, status
│   │   ├── policies.py         GET/POST /api/policies
│   │   ├── payments.py         POST /api/payment-intent
│   │   ├── uniswap.py          GET /api/token/price, POST /api/token/quote
│   │   └── wallet.py           POST /api/wallet/connect
│   ├── services/
│   │   ├── agent_service.py    AI payment intent parsing
│   │   ├── policy_service.py   Policy creation, on-chain reads
│   │   └── web3_service.py     Web3 connection, contract loading
│   └── templates/
│       └── dashboard.html      Single-page frontend
├── contracts/
│   ├── src/PolicyManager.sol   Solidity smart contract
│   └── abi/PolicyManager.json  Contract ABI
├── tests/
│   ├── test_api.py             33 integration tests
│   ├── test_validation.py      21 Pydantic schema unit tests
│   ├── test_policy_service.py  6 unit tests (mocked Web3)
│   └── test_agent_service.py   6 unit tests (mocked AI)
├── .github/workflows/
│   ├── tests.yml               pytest CI
│   └── audit.yml               Slither + Mythril smart contract audit
├── Dockerfile
├── docker-compose.yml
├── requirements.txt
└── .env.example""")

add_heading(doc, '15.2 Backend Components', level=2)

add_heading(doc, 'app.py — Entry Point', level=3)
add_paragraph(doc, 'Initialises Flask, SQLAlchemy, and Flask-Limiter. Registers all route blueprints. Creates database tables on startup. Seeds demo data if SEED_DATA=true. Exposes:')
for item in [
    'GET /api/dashboard — aggregated KPIs, transactions, activity, wallet, policies with remaining budgets',
    'GET /api/charts/payments — 7-day payment chart data',
    'GET /api/contract/status — Web3 connection status',
]:
    doc.add_paragraph(item, style='List Bullet')

add_heading(doc, 'models.py — Database Models', level=3)
add_table(doc,
    ['Model', 'Table', 'Key Fields'],
    [
        ['Policy', 'policy', 'id, agent, token, total_budget, per_tx_limit, purpose, created_at'],
        ['Transaction', 'transaction', 'id, tx_id, amount, status, policy_id, created_at'],
        ['ActivityLog', 'activity_log', 'id, action, text, time'],
        ['PaymentIntent', 'payment_intent', 'id, task, result_json, created_at'],
        ['WalletState', 'wallet_state', 'id, connected, address, network, balance'],
    ]
)

add_heading(doc, 'schemas.py — Request Validation (Pydantic v2)', level=3)
add_table(doc,
    ['Schema', 'Used by', 'Validates'],
    [
        ['PolicyCreate', 'POST /api/policies', 'agent, token, totalBudget (>0), perTxLimit, purpose (non-empty)'],
        ['PaymentIntentRequest', 'POST /api/payment-intent', 'task (non-empty string)'],
        ['SwapQuoteRequest', 'POST /api/token/quote', 'tokenIn, tokenOut (known symbols), amountUSD (>0)'],
        ['WalletConnectRequest', 'POST /api/wallet/connect', 'address (valid Ethereum address or empty)'],
    ]
)

add_heading(doc, 'services/agent_service.py — AI Intent Parser', level=3)
add_paragraph(doc, 'Fallback chain:')
for item in [
    'If SYNTH_API_KEY set → calls Anthropic (claude-haiku-4-5-20251001)',
    'Else if OPENAI_API_KEY set → calls OpenAI (gpt-4o-mini)',
    'Else → returns a demo response with a random realistic amount',
]:
    doc.add_paragraph(item, style='List Number')
add_paragraph(doc, 'Returns: {"recipient": ..., "amount": ..., "token": ..., "purpose": ...}')

add_heading(doc, 'secrets.py — Secret Loader', level=3)
add_paragraph(doc, 'Detects whether the app is running on Cloud Run (K_SERVICE env var set by the runtime):')
for item in [
    'Cloud Run: fetches each secret from GCP Secret Manager. Secrets already injected via --update-secrets are skipped.',
    'Local: calls python-dotenv load_dotenv() to read the .env file.',
]:
    doc.add_paragraph(item, style='List Bullet')

add_heading(doc, 'extensions.py — Rate Limiting', level=3)
for item in [
    'Default limits: 100 requests/day, 10 requests/minute per IP',
    'Login endpoint: 5 requests/minute, 20 requests/hour (brute force protection)',
    'Payment intent endpoint: 10 requests/minute',
    'Payment execution endpoint: 20 requests/minute',
    'Tests are exempted via @limiter.request_filter when TESTING=True',
]:
    doc.add_paragraph(item, style='List Bullet')

add_heading(doc, '15.3 Frontend', level=2)
add_paragraph(doc, 'The entire UI is a single file: app/templates/dashboard.html.')
add_paragraph(doc, 'Key design decisions:')
for item in [
    'No JavaScript framework — plain JS with fetch()',
    'All DOM manipulation uses createElement and textContent (never innerHTML) to prevent XSS',
    'Login overlay rendered on page load; removed after successful /api/auth/login call',
    'Single-page navigation: sidebar links toggle CSS display on sections, no page reloads',
    'Error toasts appear bottom-right for API failures',
    'Loading spinners on buttons during async calls',
]:
    doc.add_paragraph(item, style='List Bullet')

add_heading(doc, '15.4 Smart Contract', level=2)
add_paragraph(doc, 'File: contracts/src/PolicyManager.sol')
add_paragraph(doc, 'The PolicyManager Solidity contract enforces payment policies on the Ethereum blockchain.')
add_table(doc,
    ['Function', 'Visibility', 'Description'],
    [
        ['createPolicy(agent, token, totalBudget, perTxLimit, validFrom, validUntil, purpose)', 'onlyOwner', 'Registers a new policy on-chain'],
        ['approvePayment(policyId, amount)', 'onlyOwner', 'Checks amount against policy limits; reverts if exceeded'],
        ['getPolicy(policyId)', 'view', 'Returns policy details and remaining budget'],
    ]
)

add_heading(doc, '15.5 Tests', level=2)
add_paragraph(doc, 'Run the full test suite (requires a local PostgreSQL instance with stablepayguard_test database):')
add_code(doc, 'python -m pytest tests/ -v --cov=app --cov-report=term-missing')
add_table(doc,
    ['File', 'Tests', 'What they cover'],
    [
        ['test_api.py', '37', 'All HTTP endpoints, auth, validation, persistence, remaining budget, transaction filtering'],
        ['test_validation.py', '27', 'Pydantic schemas — PolicyCreate, PaymentIntentRequest, PaymentExecuteRequest, SwapQuoteRequest, WalletConnectRequest'],
        ['test_policy_service.py', '6', 'create_policy() in demo mode and with mocked Web3 failure'],
        ['test_agent_service.py', '6', 'generate_payment_intent() with mocked Anthropic and OpenAI'],
    ]
)

add_heading(doc, '15.6 CI/CD', level=2)
add_heading(doc, 'tests.yml', level=3)
add_paragraph(doc, 'Runs on every push or pull request touching app/, tests/, or requirements.txt:')
for item in [
    'Spins up a Postgres 16 service container',
    'Installs Python 3.11 and dependencies',
    'Runs pytest tests/ -v --cov=app',
]:
    doc.add_paragraph(item, style='List Number')

add_heading(doc, 'audit.yml', level=3)
add_paragraph(doc, 'Runs Slither and Mythril static analysis on the smart contract:')
for item in [
    'Slither — checks for common Solidity vulnerabilities (reentrancy, integer overflow, etc.)',
    'Mythril — symbolic execution to find deeper security issues',
]:
    doc.add_paragraph(item, style='List Bullet')

# ---------------------------------------------------------------------------
# Section 16
# ---------------------------------------------------------------------------

add_heading(doc, '16. Going Live on the Blockchain', level=1)
add_paragraph(doc, 'To switch from demo mode to live on-chain enforcement on the Ethereum Sepolia testnet:')

add_heading(doc, 'Step 1 — Get an RPC Endpoint', level=2)
add_paragraph(doc, 'Sign up at infura.io or alchemy.com. Create a Sepolia project and copy the RPC URL:')
add_code(doc, 'https://sepolia.infura.io/v3/YOUR_PROJECT_ID')

add_heading(doc, 'Step 2 — Prepare a Wallet', level=2)
add_paragraph(doc, 'You need an Ethereum wallet with Sepolia ETH (free from a Sepolia faucet). Export the private key. Never use a wallet holding real funds.')

add_heading(doc, 'Step 3 — Deploy the Smart Contract', level=2)
add_code(doc, 'cd contracts\nforge create src/PolicyManager.sol:PolicyManager \\\n  --rpc-url $RPC_URL \\\n  --private-key $PRIVATE_KEY')
add_paragraph(doc, 'Note the deployed contract address.')

add_heading(doc, 'Step 4 — Update Cloud Run Environment Variables', level=2)
add_code(doc, 'gcloud run services update stablepayguard \\\n  --region us-east1 \\\n  --project stablepayguard \\\n  --update-env-vars "RPC_URL=https://sepolia.infura.io/v3/YOUR_ID,PRIVATE_KEY=0xyourkey,POLICY_CONTRACT=0xDeployedAddress,OWNER_WALLET=0xYourWalletAddress"')

add_heading(doc, 'Step 5 — Verify', level=2)
add_paragraph(doc, 'Visit /api/contract/status. It should return:')
add_code(doc, '{"mode": "live", "web3_connected": true, "contract_loaded": true}')
add_paragraph(doc, 'Policy creation will now submit real Sepolia transactions, viewable at https://sepolia.etherscan.io.')

# ---------------------------------------------------------------------------
# Section 17
# ---------------------------------------------------------------------------

add_heading(doc, '17. Security', level=1)

add_heading(doc, 'Endpoint Protection', level=2)
add_table(doc,
    ['Endpoint', 'Auth Required'],
    [
        ['POST /api/auth/login', 'No'],
        ['POST /api/auth/logout', 'No'],
        ['GET /api/auth/status', 'No'],
        ['GET /api/policies', 'No'],
        ['POST /api/policies', 'Yes'],
        ['POST /api/policies/<id>/deactivate', 'Yes'],
        ['POST /api/payment-intent', 'Yes'],
        ['POST /api/payment', 'Yes'],
        ['GET /api/transactions', 'No'],
        ['POST /api/wallet/connect', 'Yes'],
        ['GET /api/dashboard', 'No'],
        ['GET /api/charts/payments', 'No'],
        ['GET /api/token/price/<symbol>', 'No'],
        ['GET /api/token/prices', 'No'],
        ['POST /api/token/quote', 'No'],
        ['GET /api/contract/status', 'No'],
        ['GET /', 'No'],
    ]
)
add_paragraph(doc, 'Authentication is session-based. The session cookie is signed with SECRET_KEY. Sessions expire when the browser closes.')

add_heading(doc, 'Rate Limiting', level=2)
for item in [
    'All endpoints: 100 requests/day, 10 requests/minute per IP',
    'Payment intent endpoint: additionally capped at 10 requests/minute',
    'Rate limiting is disabled during automated tests',
]:
    doc.add_paragraph(item, style='List Bullet')

add_heading(doc, 'Input Validation', level=2)
add_paragraph(doc, 'All state-modifying endpoints validate request bodies via Pydantic before processing. Invalid or missing fields return 400 Bad Request with field-level error details.')

add_heading(doc, 'XSS Prevention', level=2)
add_paragraph(doc, 'All user-supplied data rendered in the dashboard is inserted via textContent (not innerHTML), preventing cross-site scripting.')

add_heading(doc, 'What Should Be Hardened Before Production', level=2)
for item in [
    'Replace single-user password auth with a proper identity provider (OAuth, SSO)',
    'Enable Cloud SQL private IP (VPC) instead of the Cloud SQL Auth Proxy socket',
    'Add HTTPS-only enforcement and security headers (HSTS, CSP)',
    'Implement per-user audit logging',
    'Rotate the PRIVATE_KEY to a hardware wallet or KMS-managed key before handling real funds',
]:
    doc.add_paragraph(item, style='List Bullet')
add_paragraph(doc, 'Note: All secrets (SECRET_KEY, ADMIN_PASSWORD, API keys, PRIVATE_KEY) are already stored in Google Secret Manager — this has been implemented.')

# ---------------------------------------------------------------------------
# Section 18
# ---------------------------------------------------------------------------

add_heading(doc, '18. Troubleshooting', level=1)

add_heading(doc, '503 Service Unavailable on Cloud Run', level=2)
add_paragraph(doc, 'The container is crashing on startup. Check the logs:')
add_code(doc, 'gcloud logging read \\\n  "resource.type=cloud_run_revision AND resource.labels.service_name=stablepayguard AND severity>=ERROR" \\\n  --project stablepayguard --limit 10')
add_paragraph(doc, 'Common causes:')
for item in [
    'DATABASE_URL not set or wrong — check env vars',
    'Cloud SQL instance not running — check: gcloud sql instances describe stablepayguard-db --project stablepayguard',
    'Import error in Python code — check the full traceback in logs',
]:
    doc.add_paragraph(item, style='List Bullet')

add_heading(doc, 'Login Not Working', level=2)
for item in [
    'Verify ADMIN_PASSWORD env var is set correctly',
    'Check that SECRET_KEY is set (sessions will not persist without it)',
    'Try in a private/incognito window to rule out stale session cookies',
]:
    doc.add_paragraph(item, style='List Bullet')

add_heading(doc, 'Database Connection Refused (Local)', level=2)
add_paragraph(doc, 'Easiest fix: comment out DATABASE_URL in your .env file. The app will fall back to SQLite automatically and restart cleanly.')
for item in [
    'If you need PostgreSQL: ensure it is running (pg_isready -h localhost)',
    'Verify DATABASE_URL matches your local Postgres user/password/dbname',
    'If using Docker Compose: ensure the db service is healthy before the web service starts',
]:
    doc.add_paragraph(item, style='List Bullet')

add_heading(doc, 'Payment Intent Returns Demo Response', level=2)
for item in [
    'No AI API key is configured — set SYNTH_API_KEY or OPENAI_API_KEY',
    'Check logs for API key errors if you have set a key',
]:
    doc.add_paragraph(item, style='List Bullet')

add_heading(doc, 'Policy Creation Succeeds but Mode Stays "demo"', level=2)
for item in [
    'RPC_URL, PRIVATE_KEY, and POLICY_CONTRACT must all be set for live mode',
    'Check /api/contract/status — it shows exactly which components are connected',
]:
    doc.add_paragraph(item, style='List Bullet')

add_heading(doc, 'Tests Failing with Database Connection Error', level=2)
for item in [
    'Create the test database: createdb -U stablepayguard stablepayguard_test',
    'Ensure the DATABASE_URL in tests/test_api.py matches your local Postgres credentials',
]:
    doc.add_paragraph(item, style='List Bullet')

# ---------------------------------------------------------------------------
# Section 19
# ---------------------------------------------------------------------------

add_heading(doc, '19. Known Limitations', level=1)
add_table(doc,
    ['Limitation', 'Detail'],
    [
        ['Single user only', 'One admin account with one shared password. No user roles, no per-user policies.'],
        ['Demo AI in free tier', 'Without an Anthropic or OpenAI key, payment intent parsing returns a static demo response.'],
        ['No email or webhook alerts', 'No notification system for rejected payments or policy breaches.'],
        ['No audit log export', 'The activity feed is viewable in the UI only. No CSV/PDF export.'],
        ['No policy editing', 'Policies can be created or deactivated but not modified once created.'],
        ['db-f1-micro Cloud SQL tier', 'The deployed instance uses the smallest available tier — not suitable for production load.'],
        ['Sepolia testnet only', 'Smart contract integration targets Ethereum Sepolia. Mainnet requires redeployment.'],
        ['No MetaMask integration', 'Wallet connection accepts a manually typed address only. No WalletConnect or MetaMask.'],
        ['Single region', 'App and database are both in us-east1. No multi-region failover.'],
    ]
)

# ---------------------------------------------------------------------------
# Section 20
# ---------------------------------------------------------------------------

add_heading(doc, '20. Full API Reference', level=1)
add_paragraph(doc, 'All endpoints exposed by the Flask backend. Endpoints marked Auth=Yes require an active login session (POST /api/auth/login first).')

add_heading(doc, 'Auth', level=2)
add_table(doc,
    ['Method', 'Path', 'Auth', 'Description'],
    [
        ['POST', '/api/auth/login', 'No', 'Login. Body: {"password": "..."}. Returns {"success": true} on match.'],
        ['POST', '/api/auth/logout', 'No', 'Clears the session cookie.'],
        ['GET', '/api/auth/status', 'No', 'Returns {"authenticated": bool}.'],
    ]
)

add_heading(doc, 'Dashboard', level=2)
add_table(doc,
    ['Method', 'Path', 'Auth', 'Description'],
    [
        ['GET', '/api/dashboard', 'No', 'KPIs, recent transactions, activity feed, wallet state, policies with remaining budgets.'],
        ['GET', '/api/charts/payments', 'No', 'Last 7 days of completed payment volume grouped by day of week.'],
        ['GET', '/api/contract/status', 'No', 'Web3 connection and contract load status. Returns mode: live or demo.'],
    ]
)

add_heading(doc, 'Policies', level=2)
add_table(doc,
    ['Method', 'Path', 'Auth', 'Description'],
    [
        ['GET', '/api/policies', 'No', 'List all policies.'],
        ['POST', '/api/policies', 'Yes', 'Create policy. Body: agent, token, totalBudget, perTxLimit, validFrom, validUntil, purpose.'],
        ['GET', '/api/policies/<id>', 'No', 'Policy detail — live from contract if connected, DB otherwise.'],
        ['POST', '/api/policies/<id>/deactivate', 'Yes', 'Deactivate policy on-chain (or demo mode fallback). Permanent.'],
    ]
)

add_heading(doc, 'Payments and Transactions', level=2)
add_table(doc,
    ['Method', 'Path', 'Auth', 'Description'],
    [
        ['POST', '/api/payment-intent', 'Yes', 'Parse plain-English task into structured payment JSON via AI. Rate limit: 10/min.'],
        ['POST', '/api/payment', 'Yes', 'Execute payment. Body: policy_id, recipient, amount, purpose. Validates budget, submits on-chain or demo. Rate limit: 20/min.'],
        ['GET', '/api/transactions', 'No', 'Filterable ledger. Query params: ?status=, ?policy=, ?limit= (max 200), ?offset=.'],
    ]
)

add_heading(doc, 'Tokens', level=2)
add_table(doc,
    ['Method', 'Path', 'Auth', 'Description'],
    [
        ['GET', '/api/token/price/<symbol>', 'No', 'Live USD price for ETH, USDC, DAI, USDT, WBTC. Results cached 60 seconds.'],
        ['POST', '/api/token/quote', 'No', 'Swap estimate. Body: tokenIn, tokenOut, amountUSD.'],
        ['GET', '/api/token/prices', 'No', 'All five token prices in one call.'],
    ]
)

add_heading(doc, 'Wallet', level=2)
add_table(doc,
    ['Method', 'Path', 'Auth', 'Description'],
    [
        ['POST', '/api/wallet/connect', 'Yes', 'Connect wallet. Body: {"address": "0x..."}. Omit address to use OWNER_WALLET from env.'],
    ]
)

# ---------------------------------------------------------------------------
# Section 21
# ---------------------------------------------------------------------------

add_heading(doc, '21. Architecture Deep Dive', level=1)

add_heading(doc, '21.1 System Diagram', level=2)
add_code(doc, """\
Browser (HTML / CSS / JS)
        |  REST API (JSON)
        v
Flask Backend  --  Gunicorn (Cloud Run) / dev server (local)
  |-- 5 route blueprints (auth, policies, payments, uniswap, wallet)
  |-- 4 service modules  (agent, policy, uniswap, web3)
  |-- Pydantic validation + Flask-Limiter rate limiting
  `-- secrets.py  ->  GCP Secret Manager (Cloud Run) / .env (local)
        |
        |-- PostgreSQL  --  Cloud SQL (GCP) / SQLite (local fallback)
        |
        |-- Ethereum Sepolia  --  PolicyManager.sol
        |       via Infura RPC
        |
        `-- Uniswap v3 Subgraph  --  live token prices (60s cache)""")

add_heading(doc, '21.2 Component Details', level=2)

add_heading(doc, 'app/app.py — Entry Point', level=3)
add_paragraph(doc, 'Responsibilities:')
for item in [
    'Call load_secrets() to populate os.environ from GCP Secret Manager (Cloud Run) or .env (local)',
    'Configure SQLAlchemy with DATABASE_URL — falls back to SQLite if not set',
    'Register 5 blueprints',
    'Run db.create_all() and seed demo data on first run',
    'Serve dashboard page (GET /)',
    'Expose dashboard API, chart API, and contract status API',
]:
    doc.add_paragraph(item, style='List Bullet')

add_heading(doc, 'app/secrets.py — Secret Loader', level=3)
add_paragraph(doc, 'Detects environment via K_SERVICE env var (set automatically by Cloud Run):')
for item in [
    'Cloud Run: fetches secrets from GCP Secret Manager for any not already injected via --update-secrets',
    'Local: calls python-dotenv load_dotenv() to read .env file',
]:
    doc.add_paragraph(item, style='List Bullet')

add_heading(doc, 'app/models.py — ORM Models', level=3)
add_paragraph(doc, 'Five SQLAlchemy models: Policy, Transaction, ActivityLog, PaymentIntent, WalletState.')
for item in [
    'Transaction.policy is a foreign key to policies.id (ondelete=SET NULL)',
    'Transaction.status has a CHECK constraint: only Completed, Pending, Declined accepted',
    'created_at columns on Transaction, ActivityLog, PaymentIntent are indexed for query performance',
    'SQLite FK enforcement enabled via PRAGMA foreign_keys=ON event listener',
]:
    doc.add_paragraph(item, style='List Bullet')

add_heading(doc, 'app/services/agent_service.py — AI Intent Parser', level=3)
add_paragraph(doc, 'Fallback chain:')
for item in [
    'SYNTH_API_KEY set -> Anthropic (claude-haiku-4-5-20251001)',
    'OPENAI_API_KEY set -> OpenAI (gpt-4o-mini)',
    'Neither set -> demo mode with random realistic amount ($100-$5000)',
]:
    doc.add_paragraph(item, style='List Number')
add_paragraph(doc, 'Validates parsed JSON has recipient, amount (numeric), and purpose fields before returning.')

add_heading(doc, 'app/services/uniswap_service.py — Token Prices', level=3)
for item in [
    'Queries Uniswap v3 Subgraph (The Graph) via GraphQL',
    'Stablecoins (USDC, DAI, USDT) hardcoded to $1.00',
    'ETH and WBTC calculated from pool derivedETH x ethPriceUSD',
    '60-second in-memory TTL cache per symbol to avoid repeated Subgraph calls',
]:
    doc.add_paragraph(item, style='List Bullet')

add_heading(doc, 'app/services/web3_service.py — Blockchain Interface', level=3)
for item in [
    'Initialises Web3.HTTPProvider(RPC_URL) on startup',
    'Loads PolicyManager.json ABI from contracts/abi/',
    'approve_payment(): builds, signs, and sends approvePayment transaction using PRIVATE_KEY',
    'Falls back gracefully to demo mode if RPC_URL or POLICY_CONTRACT are not set',
]:
    doc.add_paragraph(item, style='List Bullet')

add_heading(doc, '21.3 Database Schema', level=2)
add_table(doc,
    ['Table', 'Key Columns', 'Notes'],
    [
        ['policies', 'id (PK), agent, token, budget, purpose, tx_hash, created_at', 'Indexed on created_at'],
        ['transactions', 'id (PK), recipient, policy (FK), amount, status, hash, created_at', 'FK -> policies.id; status CHECK constraint; indexed on policy, created_at'],
        ['activity_logs', 'id, action, text, time, created_at', 'Indexed on created_at'],
        ['payment_intents', 'id, task, recipient, amount, purpose, mode, error, created_at', 'Indexed on created_at'],
        ['wallet_state', 'id (singleton=1), connected, address', 'Single row, updated in place'],
    ]
)

add_heading(doc, '21.4 Data Flow: Payment Execution', level=2)
add_code(doc, """\
POST /api/payment  {policy_id, recipient, amount, purpose}
        |
        v
  Pydantic validation (PaymentExecuteRequest)
        |
        v
  DB: policy exists?  --No--> 404
        | Yes
        v
  DB: remaining budget >= amount?  --No--> 422 + activity log
        | Yes
        v
  web3_service.approve_payment()
    |-- Connected: on-chain tx, returns real 64-char hash
    `-- Not connected: skip (demo hash generated by store)
        |
        v
  store.new_tx() -> persist transaction record
        |
        v
  store.add_activity() -> activity log entry
        |
        v
  201 {tx_id, hash, mode: "live"|"demo"}""")

add_heading(doc, '21.5 Data Flow: Secret Loading', level=2)
add_code(doc, """\
App starts (app.py calls load_secrets())
        |
        |-- K_SERVICE set? (Cloud Run runtime)
        |     |  Yes
        |     `-> GCP Secret Manager
        |           fetch each secret not already in os.environ
        |           (--update-secrets may have already injected some)
        |
        `-- K_SERVICE not set? (local dev)
              `-> load_dotenv() reads .env file
        |
        v
  os.environ populated -> app initialisation continues""")

add_heading(doc, '21.6 Deployment Architecture', level=2)
add_code(doc, """\
Developer machine
        |  gcloud run deploy --source .
        v
Cloud Build
        |  builds Docker image from Dockerfile
        v
Artifact Registry
        |  stores image
        v
Cloud Run (us-east1)  service: stablepayguard
  Secrets: mounted from Secret Manager
  DB:      Cloud SQL PostgreSQL via DATABASE_URL env var
  Port:    8080

Gunicorn command (Dockerfile):
  gunicorn --bind 0.0.0.0:$PORT --workers 1 --threads 8
           --timeout 60 --chdir /app/app app:app""")

add_paragraph(doc, 'Project structure in the repository:')
add_code(doc, """\
StablePayGuard/
  app/
    app.py                   Flask entry point, dashboard + chart routes
    models.py                SQLAlchemy ORM models
    schemas.py               Pydantic request validation schemas
    store.py                 DB helper functions and seed data
    extensions.py            Flask-Limiter setup
    utils.py                 login_required decorator, validate_request helper
    secrets.py               Secret loader: GCP Secret Manager / .env fallback
    routes/
      auth.py                Login / logout / status
      policies.py            Policy CRUD + deactivation
      payments.py            Payment intent, execution, transaction listing
      uniswap.py             Token price and swap quote
      wallet.py              Wallet connection
    services/
      agent_service.py       AI payment intent parsing
      policy_service.py      On-chain policy creation and deactivation
      uniswap_service.py     Uniswap v3 Subgraph price feeds with 60s cache
      web3_service.py        Web3 connection and contract interface
    templates/
      dashboard.html         Single-page dashboard UI
  contracts/
    src/PolicyManager.sol    Solidity smart contract (Sepolia)
    abi/PolicyManager.json   Contract ABI
    audit/slither_report.md  Slither security audit results
    deployment.json          Deployed contract addresses
  tests/
    test_api.py              Integration tests (requires PostgreSQL)
    test_agent_service.py    AI service unit tests
    test_policy_service.py   Policy service unit tests
    test_validation.py       Schema validation unit tests
  scripts/deploy.py          Contract deployment script
  .github/workflows/
    tests.yml                CI: pytest on every push
    audit.yml                CI: Slither on contract changes
  .env.example               Environment variable template
  Dockerfile                 Cloud Run container
  docker-compose.yml         Local full-stack: app + PostgreSQL
  requirements.txt""")

# ---------------------------------------------------------------------------
# Section 22
# ---------------------------------------------------------------------------

add_heading(doc, '22. Roadmap', level=1)
add_table(doc,
    ['Phase', 'Status', 'Scope'],
    [
        ['1 — Core Platform', 'Complete', 'Policy engine, AI parsing, dashboard, on-chain enforcement, security audit'],
        ['2 — Multi-user Auth', 'Planned', 'JWT tokens, role-based access, multi-user dashboards'],
        ['3 — Real Payment Rails', 'Planned', 'USDC transfers, bank APIs, escrow mechanisms'],
        ['4 — AI Agent Platform', 'Planned', 'Anomaly detection, spending predictions, autonomous agent orchestration'],
    ]
)

# ---------------------------------------------------------------------------
# Section 23
# ---------------------------------------------------------------------------

add_heading(doc, '23. Business Scenario', level=1)

add_heading(doc, '23.1 Ideal Customer Profile', level=2)
add_paragraph(doc, 'The highest-value customer for this platform is a large global enterprise paying vendors, contractors, or partners across multiple countries at high volume.')
add_table(doc,
    ['Signal', 'Why It Matters'],
    [
        ['Cross-border payment volume', 'SWIFT wires cost $35–$75 each and take 1–5 days. USDC on Base costs $0.02 and settles in 2 seconds. At 500 international payments/month the fee savings alone are $17,500–$37,500/month.'],
        ['High invoice volume', 'Human review of every invoice costs $15–$50 in staff time. At 1,000 invoices/month that is $15,000–$50,000/month in AP labor that automation eliminates for routine transactions.'],
        ['Audit and compliance requirements', 'Blockchain audit trail is immutable and independently verifiable without internal system access — critical for SOX compliance, external audit, and fraud investigation.'],
        ['Crypto treasury or digital asset strategy', 'Increasingly common among Fortune 500 companies. Paying vendors in USDC avoids FX conversion entirely when the treasury already holds stablecoins.'],
        ['Globally distributed workforce', 'Contractor in Singapore, vendor in Germany, supplier in Brazil — one payment rail, same cost, same speed, regardless of destination.'],
    ]
)
add_paragraph(doc, 'Industries With the Strongest Fit:')
add_table(doc,
    ['Industry', 'Specific Pain Point Solved'],
    [
        ['Technology (large enterprise)', 'Global SaaS vendor payments, remote contractor disbursements'],
        ['Manufacturing', 'Cross-border supplier payments, multi-currency invoice settlement'],
        ['Healthcare', 'International lab vendor payments, contractor reimbursements across jurisdictions'],
        ['Real Estate', 'Cross-border property management, international contractor payments'],
        ['Retail / E-commerce', 'Global fulfillment partner settlements, international marketplace fees'],
        ['Professional Services', 'Cross-border subcontractor payments, international expert fees'],
        ['Non-Profit / NGO', 'Grant disbursements to international programs with full audit trail'],
        ['Government / Public Sector', 'Vendor payments within procurement rules, independently auditable'],
    ]
)
add_paragraph(doc, 'The common thread: any large organization making repetitive cross-border payments within known rules, where the cost of human approval and wire transfer friction is significant.')

add_heading(doc, '23.2 The Core Problem This Solves', level=2)
add_paragraph(doc, 'Traditional Accounts Payable works like this:')
for item in [
    '1. Invoice arrives',
    '2. Human reviews it',
    '3. Human checks budget',
    '4. Human checks vendor approval',
    '5. Human approves payment',
    '6. Payment executes',
    '7. Human files the record',
]:
    doc.add_paragraph(item, style='List Bullet')
add_paragraph(doc, 'This process costs time, money, and introduces human error — for every single invoice, including the routine $1,800 AWS bill that has arrived on the same date every month for three years.')
q = add_paragraph(doc, 'Humans define the rules once. AI agents execute within those rules continuously. The smart contract enforces the rules autonomously. Humans only act on exceptions.')
q.runs[0].bold = True

add_heading(doc, '23.3 Why This Needs Crypto', level=2)
add_paragraph(doc, 'The enforcement logic can technically be built without crypto. Tools like Coupa, Tipalti, and Bill.com do something similar with traditional databases.')
add_table(doc,
    ['Capability', 'Traditional Database', 'Blockchain Smart Contract'],
    [
        ['Audit trail integrity', 'Internal database — can be edited by a DBA, wiped in a breach, or altered by an insider', 'Immutable — no one, including the company itself, can alter a confirmed transaction'],
        ['Enforcement authority', 'Runs on servers the company controls — a compromised server or rogue admin can bypass rules', 'Smart contract cannot be bypassed even by the people who wrote it'],
        ['Cross-border settlement', 'SWIFT wire: $35–$75, 1–5 days, multiple intermediaries', 'USDC on Base: $0.02, 2 seconds, no intermediaries'],
        ['Independent verifiability', 'Requires internal system access for audit', 'Anyone with the contract address can verify every transaction'],
        ['Fraud surface', 'Any human in the approval chain, any server in the payment path', 'Only the policy owner wallet can change rules'],
    ]
)
add_paragraph(doc, 'When crypto is not necessary: A purely domestic company paying five local vendors in USD, with no cross-border volume, no external audit requirement, and simple AP needs does not need this platform.')
add_paragraph(doc, 'When crypto is clearly the right choice:')
for item in [
    'Cross-border payment volume where wire fees are a real cost',
    'Industries with external audit or compliance requirements where audit trail integrity must be beyond internal question',
    'Companies already holding crypto or stablecoin treasuries',
    'Organizations paying global contractors where banking access varies',
]:
    doc.add_paragraph(item, style='List Bullet')

add_heading(doc, '23.4 Cross-Border Payment Economics', level=2)
add_paragraph(doc, '$1,000 payment to a contractor in Singapore:')
add_table(doc,
    ['Method', 'Sender Fee', 'Receiver Fee', 'FX Cost', 'Settlement Time'],
    [
        ['SWIFT wire', '$25–$50', '$10–$25', '1–3% spread', '1–5 business days'],
        ['PayPal International', '3–5% ($30–$50)', 'None', '2–4% spread', 'Minutes to days'],
        ['USDC on Base', '~$0.02', 'None', 'None (stablecoin)', '~2 seconds'],
    ]
)
add_paragraph(doc, 'At 500 international payments per month, SWIFT costs $17,500–$37,500 in fees alone. USDC on Base costs approximately $10.')
add_paragraph(doc, 'Why Base is the right network for enterprise use:')
add_table(doc,
    ['Property', 'Detail'],
    [
        ['Cost per transfer', '~$0.01–$0.05 regardless of amount sent'],
        ['Settlement time', '~2 seconds (practical finality)'],
        ['Full finality', '~1 minute (statistically irreversible)'],
        ['USDC type', 'Native — issued directly by Circle, not bridged'],
        ['Security model', 'Ethereum L2 — inherits Ethereum security'],
        ['Institutional backing', 'Operated by Coinbase — compliance-friendly'],
        ['Regulatory clarity', 'USDC issued under US money transmission licenses'],
    ]
)

add_heading(doc, '23.5 Agent A: SaaS Vendor Payments', level=2)
add_paragraph(doc, 'A global enterprise pays three SaaS vendors every month: AWS, GitHub, and Datadog. Rules: total monthly budget $15,000; maximum per transaction $2,000; approved vendors only; payments only within the current calendar month.')
add_paragraph(doc, 'PHASE 1 — One-Time Human Setup:')
add_table(doc,
    ['Step', 'Who', 'What They Do', 'Time'],
    [
        ['1', 'CFO / AP Manager', 'Decides which vendors are approved: AWS, GitHub, Datadog', '10 min'],
        ['2', 'CFO / AP Manager', 'Records each vendor\'s USDC payment wallet address', '10 min'],
        ['3', 'CFO / AP Manager', 'Generates a cryptographic hash of the approved vendor list stored on-chain', '2 min'],
        ['4', 'CFO / AP Manager', 'Creates the policy: agent wallet, token, $15,000 budget, $2,000 per-tx limit, valid date window, purpose hash', '2 min'],
        ['5', 'CFO / AP Manager', 'Clicks Deploy — policy written to the blockchain permanently', '1 min'],
        ['6', 'IT / Dev', 'Configures Agent A to monitor the AP invoice inbox', '1 hour'],
        ['7', 'Vendors', 'Already send structured invoices — no change needed', '0 min'],
    ]
)
add_paragraph(doc, 'Total human setup time: approximately 1.5 hours, done once.')
add_paragraph(doc, 'PHASE 2 — Ongoing Agent Operation (No Human Required). When an invoice arrives from AWS for $1,800:')
add_table(doc,
    ['Step', 'Who', 'What Happens'],
    [
        ['1', 'Agent A', 'Detects new invoice in the AP inbox'],
        ['2', 'Agent A', 'Parses vendor name, amount, due date, payment wallet address'],
        ['3', 'Agent A', 'Checks: is this vendor on the approved list? Does the wallet address match?'],
        ['4', 'Agent A', 'Calls the PolicyManager smart contract: "Can I pay $1,800 against policy POL-101?"'],
        ['5', 'Smart Contract', 'Checks: amount <= per-tx limit ($1,800 <= $2,000)'],
        ['6', 'Smart Contract', 'Checks: running total + amount <= monthly budget ($12,400 + $1,800 <= $15,000)'],
        ['7', 'Smart Contract', 'Checks: today is within the valid date window'],
        ['8', 'Smart Contract', 'Checks: vendor hash matches approved vendor list'],
        ['9', 'Smart Contract', 'Approves the payment, updates running spend total on-chain'],
        ['10', 'Agent A', 'Executes USDC transfer to AWS wallet — settles in ~2 seconds'],
        ['11', 'Smart Contract', 'Emits PaymentApproved event — permanently recorded on-chain'],
        ['12', 'Dashboard', 'Transaction appears in real time: amount, vendor, policy ID, network hash'],
        ['13', 'CFO', 'Sees it on the dashboard — no action needed'],
    ]
)
add_paragraph(doc, 'Total human time for this payment: zero minutes.')
add_paragraph(doc, 'PHASE 3 — Exception Handling:')
add_table(doc,
    ['Situation', 'What the Contract Does', 'What the Agent Does', 'Human Action Required'],
    [
        ['Invoice for $1,800 from AWS (normal)', 'Approves', 'Pays', 'None'],
        ['Invoice for $2,400 from AWS (over per-tx limit)', 'Rejects', 'Flags for human review', 'CFO decides: approve as exception or reject'],
        ['Invoice from unknown vendor', 'Rejects — not on approved list', 'Flags for human review', 'AP Manager investigates'],
        ['Monthly budget reaches $14,800 — $600 invoice arrives', 'Rejects — would exceed $15,000', 'Flags for human review', 'CFO decides: increase budget or defer to next month'],
        ['Invoice arrives after policy expiry date', 'Rejects — policy expired', 'Flags for human review', 'CFO renews policy for new month'],
        ['Duplicate invoice submitted twice', 'Rejects — budget already debited', 'Flags as duplicate', 'AP Manager confirms and closes'],
    ]
)

add_heading(doc, '23.6 Agent B: Global Contractor Payments', level=2)
add_paragraph(doc, 'A global enterprise pays independent contractors across multiple countries every two weeks. Rules: total monthly budget $50,000; maximum per transaction $5,000; pre-approved contractors only; payments only during business hours (Mon–Fri, 9am–6pm UTC).')
add_paragraph(doc, 'PHASE 1 — One-Time Human Setup:')
add_table(doc,
    ['Step', 'Who', 'What They Do', 'Time'],
    [
        ['1', 'Legal / Procurement', 'Approves contractor roster — reviews contracts, verifies identities', '1–2 days (existing onboarding)'],
        ['2', 'Procurement / AP Manager', 'Records each contractor\'s USDC wallet address — verified directly with contractor', '30 min per contractor'],
        ['3', 'Legal', 'Finalises the master contractor agreement document', 'Existing process'],
        ['4', 'AP Manager', 'Generates a cryptographic hash of the approved contractor list and master agreement', '5 min'],
        ['5', 'AP Manager', 'Creates Agent B\'s policy: $50,000 monthly budget, $5,000 per-tx limit, USDC, purpose hash', '5 min'],
        ['6', 'AP Manager', 'Sets business hours enforcement via validFrom/validUntil, renewed each week', '2 min per week'],
        ['7', 'AP Manager', 'Clicks Deploy — policy written to blockchain', '1 min'],
        ['8', 'IT / Dev', 'Configures Agent B to monitor the contractor invoice inbox', '1–2 hours'],
        ['9', 'Contractors', 'Notified of USDC payment setup — provide wallet address, receive test transaction', '15 min per contractor'],
    ]
)
add_paragraph(doc, 'PHASE 2 — Ongoing Agent Operation. When a contractor in Singapore submits an invoice for $3,200:')
add_table(doc,
    ['Step', 'Who', 'What Happens'],
    [
        ['1', 'Agent B', 'Detects new invoice — contractor name, amount, wallet address, invoice reference'],
        ['2', 'Agent B', 'Verifies: is this contractor on the approved roster?'],
        ['3', 'Agent B', 'Verifies: does the invoice reference the current active contractor agreement?'],
        ['4', 'Agent B', 'Checks current time: is it Mon–Fri, 9am–6pm UTC?'],
        ['5–9', 'Smart Contract', 'Checks: $3,200 <= $5,000 per-tx; running total <= $50,000; within valid window; hash matches'],
        ['10', 'Smart Contract', 'Approves payment, updates running spend on-chain'],
        ['11', 'Agent B', 'Executes USDC transfer — no bank, no SWIFT, no correspondent fees'],
        ['12', 'Contractor', 'Receives $3,200 USDC in Singapore in ~2 seconds'],
        ['13', 'Smart Contract', 'Emits PaymentApproved event — permanently recorded'],
    ]
)
add_paragraph(doc, 'Traditional alternative: SWIFT wire to Singapore bank, $35–$50 sender fee, $15–$25 receiver fee, 2–4 business days.')
add_paragraph(doc, 'With Agent B: $0.02, 2 seconds, full amount delivered, independently verifiable.')

add_heading(doc, '23.7 Agent C: Employee T&E Reimbursements', level=2)
add_paragraph(doc, 'Employees submit travel and expense claims after business trips. Rules: maximum per claim $500 (above this goes to manager approval); claims must match approved expense categories; processed within 24 hours.')
add_paragraph(doc, 'PHASE 2 — When an employee submits a $240 claim for client dinner receipts:')
add_table(doc,
    ['Step', 'Who', 'What Happens'],
    [
        ['1', 'Employee', 'Submits expense claim: category (client entertainment), amount ($240), receipts, date'],
        ['2', 'Agent C', 'Receives the claim via API from the expense portal'],
        ['3–5', 'Agent C', 'Checks: category in approved T&E policy? Amount within per-category limit? Current policy hash?'],
        ['6–10', 'Smart Contract', 'Checks: $240 <= $500 per-tx; running total within budget; within valid window; purposeHash matches'],
        ['11', 'Smart Contract', 'Approves payment, updates running spend on-chain'],
        ['12', 'Agent C', 'Executes USDC transfer to employee\'s registered wallet'],
        ['13', 'Employee', 'Receives $240 USDC within seconds of submission — not 2 weeks'],
    ]
)

add_heading(doc, '23.8 All Three Agents Together', level=2)
add_paragraph(doc, 'At any point, the CFO sees on the dashboard:')
add_table(doc,
    ['Agent', 'Role', 'Monthly Budget', 'Spent', 'Remaining', 'Transactions'],
    [
        ['Agent A', 'SaaS vendor payments', '$15,000', '$12,400', '$2,600', '8 completed'],
        ['Agent B', 'Contractor payments', '$50,000', '$31,200', '$18,800', '14 completed, 1 pending'],
        ['Agent C', 'T&E reimbursements', '$25,000', '$8,750', '$16,250', '37 completed, 2 flagged'],
    ]
)
add_paragraph(doc, 'No AP clerk had to touch any of the 59 completed transactions. A human was only needed for the 3 flagged exceptions.')

add_heading(doc, '23.9 Traditional AP vs. StablePayGuard', level=2)
add_table(doc,
    ['Capability', 'Traditional AP', 'StablePayGuard'],
    [
        ['Payment approval', 'Human reviews every invoice', 'Smart contract enforces rules automatically'],
        ['Audit trail', 'Internal database — can be edited', 'Immutable blockchain record — cannot be altered'],
        ['Budget enforcement', 'Policy document + manual check', 'Hard-coded in smart contract — mathematically enforced'],
        ['Vendor verification', 'Spreadsheet or ERP lookup', 'Cryptographic hash comparison on-chain'],
        ['Exception handling', 'Everything is an exception', 'Only true exceptions reach a human'],
        ['Cross-border speed', '1–5 business days', '~2 seconds'],
        ['Cross-border cost', '$35–$75 per wire', '~$0.02 per transaction'],
        ['Processing cost', '$15–$50 per invoice in staff time', 'Near zero for routine payments'],
        ['Auditability', 'Requires internal system access', 'Anyone with the contract address can verify'],
        ['Fraud surface', 'Any human in the approval chain', 'Only the policy owner wallet can change rules'],
    ]
)

add_heading(doc, '23.10 Security Considerations', level=2)
add_paragraph(doc, 'The statement "nobody can authorize a payment that violates the policy" holds true as long as the smart contract code has no bugs and the owner\'s private key is not compromised.')
add_table(doc,
    ['Risk', 'What Could Happen', 'Mitigation'],
    [
        ['Bug in smart contract', 'A logic error could allow payments that should be rejected', 'Professional audit by firms like OpenZeppelin or Certik before production'],
        ['Private key compromise', 'Attacker becomes owner and can change policies', 'Hardware wallet (Ledger/Trezor) or multi-signature requiring 2-of-3 approvals'],
        ['Compromised agent wallet', 'Attacker calls approvePayment within policy limits', 'Policy limits contain the blast radius — attacker cannot exceed what the policy allows'],
        ['Reentrancy attack', 'Malicious contract loops back during payment execution', 'Use OpenZeppelin\'s ReentrancyGuard in production contract'],
    ]
)

add_heading(doc, '23.11 The Fundamental Principle', level=2)
add_paragraph(doc, 'The smart contract is not a convenience — it is the enforcement layer.')
add_paragraph(doc, 'The agent does not need to be trusted. The CFO does not need to monitor every payment. The AP team does not need to touch routine invoices.')
add_paragraph(doc, 'The rules are on-chain. The contract enforces them. A rogue employee cannot override them. A compromised server cannot bypass them. Even the company that deployed the contract cannot silently alter a payment record after the fact.')
p = add_paragraph(doc, 'That is what makes this fundamentally different from every existing AP automation tool — and why the blockchain layer is not optional.')
p.runs[0].bold = True

add_heading(doc, '23.12 Production Roadmap', level=2)
add_table(doc,
    ['Phase', 'What Gets Added', 'Business Benefit'],
    [
        ['Phase 1 — Current', 'Policy engine, on-chain enforcement, dashboard, Uniswap pricing', 'Working prototype demonstrating the full concept'],
        ['Phase 2 — Data Layer', 'PostgreSQL persistence, user authentication, multi-user dashboards, role-based access', 'CFO, AP Manager, and auditor each see their relevant view'],
        ['Phase 3 — Real Payment Rails', 'Mainnet deployment, real USDC settlement, bank API integration (ACH/SWIFT)', 'Actual money movement at production scale'],
        ['Phase 4 — AI Agent Platform', 'Anomaly detection, spending predictions, policy auto-optimization, multi-agent coordination', 'Proactive controls — system flags unusual patterns before they become problems'],
    ]
)

# ---------------------------------------------------------------------------
# Section 24
# ---------------------------------------------------------------------------

add_heading(doc, '24. Contributing Guidelines', level=1)

add_heading(doc, 'Workflow', level=2)
for item in [
    '1. Fork the repository',
    '2. Create a feature branch: git checkout -b feature/my-feature',
    '3. Make your changes',
    '4. Commit using a clear message: git commit -m "Add feature"',
    '5. Push your branch: git push origin feature/my-feature',
    '6. Open a pull request',
]:
    doc.add_paragraph(item, style='List Bullet')

add_heading(doc, 'Guidelines', level=2)
for item in [
    'Follow PEP 8 for Python code',
    'Keep HTML/CSS changes visually consistent with the dashboard theme',
    'Document new API endpoints in the Manual',
    'Add tests where practical',
    'Keep commits focused and readable',
]:
    doc.add_paragraph(item, style='List Bullet')

add_heading(doc, 'Pull Request Checklist', level=2)
for item in [
    'Code runs locally',
    'No secrets committed',
    'Manual updated if behavior changed',
    'New configuration documented',
    'Formatting checked',
]:
    doc.add_paragraph(item, style='List Bullet')

# ---------------------------------------------------------------------------
# Save
# ---------------------------------------------------------------------------

output_path = 'StablePayGuard_Manual.docx'
doc.save(output_path)
print(f'Saved: {output_path}')
